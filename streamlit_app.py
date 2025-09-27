from __future__ import annotations

# === Standard Library ===
import math
import os
import json
import re
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime, timedelta
from io import BytesIO
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

# === Third-party ===
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # headless backend for Streamlit/servers
import matplotlib.pyplot as plt
import pyproj
from pyproj import Transformer
from shapely.geometry import Point, shape
import contextily as cx
import streamlit as st

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.text.paragraph import Paragraph

from docxtpl import DocxTemplate, InlineImage

# --- Coordinate sanity window (Gaza / Israel vicinity) ---
LAT_MIN, LAT_MAX = 29.0, 34.9
LON_MIN, LON_MAX = 33.0, 36.8

# Treat these as empty/missing
_MISSING_TOKENS = {"", "-", "—", "n/a", "na", "none", "לא ידוע", "אין", "מחק", "delete"}

# === Coords transformer (top-level init to avoid re-creating) ===
TRANSFORMER_WGS84_TO_UTM36 = pyproj.Transformer.from_crs(
    "EPSG:4326", "EPSG:32636", always_xy=True
)

# === Constants ===
TEMPLATE_FILE = "coordinations_template.docx"
HEBREW_LETTERS: List[str] = [
    "א","ב","ג","ד","ה","ו","ז","ח","ט","י","יא","יב","יג","יד","טו","טז",
    "יז","יח","יט","כ","כא","כב","כג","כד","כה","כו","כז","כח","כט","ל"
]
BIDI_CTRL_RE = re.compile(r"[\u200e\u200f\u202a-\u202e\u2066-\u2069]")
HEBREW_RANGE_RE = re.compile(r"[\u0590-\u05FF]")

PERSONNEL_TRANSLATIONS: Dict[str, str] = {
    "driver": "נהג",
    "assistant driver": "נהג משנה",
    "passenger": "נוסע",
    "pax": "נוסע",
    "patient": "מטופל",
    "patient's companion": "מלווה",
    "security escort": "נוסע - מאבטח",
    "suno": "בכיר",
    "mission team leader": "נוסע",
    "south to north": "\u200F)מדרום לצפון(\u200F",
    "north to south": "\u200F)מצפון לדרום(\u200F",
    "leaving gaza": "\u200F)יוצא מרצע(\u200F",
    "entering gaza": "\u200F)נכנס לרצע(\u200F",
}

SUNO_TITLES: Dict[str, str] = {
    "SUNC80482D": "ר' OCHA ישראל",
    "SUNJ14406D": "ר' OCHA ישראל",
    "SUNJ16507": "סגן ר' OCHA ישראל",
    "SUNB28399": "סגן ר' OCHA עזה",
    "SUNJ17294": "יועצת מיוחדת של OCHA",
    "SUNR32190": "סגנית המתאמת ההומניטרית",
    "SUNR23865D": "מ\"מ המתאם ההומניטרי",
    "AUNR72415": "ר' WHO ישראל",
    "AUNR73490D": "ר' WFP ישראל",
    "AUNB54993D": "ר' WFP עזה",
    "124500670": "סגן ר' WFP עזה",
    "SUNR33194D": "ר' UNICEF ישראל",
    "SUNJ10442D": "ר' UNICEF עזה",
    "144666595": "ר' UNMAS עזה",
    "SUNC80122D": "ר' UNDP עזה",
    "SUNJ10610D": "סגן ר' UNDP ישראל",
    "SUNC80167D": "ר' UNDSS ישראל",
    "SUNR33736D": "סגן ר' UNDSS ישראל",
    "X0N73G02": "ר' צלא עזה",
    "PB3553006": "סגן ר' צלא ישראל",
    "567985289": "מייסד ארגון WCK",
    "548965410": "ר' ANERA עזה",
}

VEHICLE_TRANSLATIONS: Dict[str, str] = {
    "truck": "משאית",
    "box truck": "משאית סגורה",
    "flatbed": "משאית פתוחה",
    "bus": "אוטובוס",
    "un": "טויוטה לנד קרוזר",
    "lc": "טויטה לנד קרוזר",
    "toyota land cruiser": "טויוטה לנד קרוזר",
    "land cruiser": "טויוטה לנד קרוזר",
}

SITE_NAME_TRANSLATIONS: Dict[str, str] = {
    "kerem shalom": "כרמש",
    "kerem shaloum": "כרמש",
    "karem salim": "כרמש",
    "karem salem": "כרמש",
    "ks": "כרמש",
    "kissufim": "מעבר 147",
    "kissuffim": "מעבר 147",
    "kisuffim": "מעבר 147",
    "ksf": "מעבר 147",
    "gate 96": "שער 96",
    "netzarim/gate96": "שער 96",
    "airport road": "פולאט",
}

END_SECTION_TOKENS: List[str] = [
    "first name",
    "last name",
    "mission of the specific personnel",
    "vehicle type",
    "plate number",
    "equipment",
    "personnel",
]

# === Data models ===
@dataclass
class Site:
    name: str
    lat: float
    lon: float

@dataclass
class PersonRow:
    name: str
    id: str
    notes: str  # role text / mission text

@dataclass
class VehicleRow:
    vehicle_type: str
    plate_number: str

# === Small helpers ===

def hebrew_index(idx_zero_based: int) -> str:
    """Return a Hebrew index label; falls back to 1-based numerals if we run out."""
    if idx_zero_based < len(HEBREW_LETTERS):
        return HEBREW_LETTERS[idx_zero_based]
    return str(idx_zero_based + 1)

def _hebrew(s: str) -> bool:
    return bool(HEBREW_RANGE_RE.search(str(s or "")))

def _sanitize_text(s: Any) -> str:
    if s is None:
        return ""
    return re.sub(r"\s+", " ", str(s)).strip()

def _sanitize_coord_text(s: Any) -> str:
    """Remove hidden bidi/control chars, normalise punctuation/spacing for coordinates."""
    s = _sanitize_text(s)
    s = BIDI_CTRL_RE.sub("", s)  # strip hidden RTL/LTR marks
    s = s.replace("،", ",").replace("‎", "").replace("‏", "")
    return s

def _norm_id_for_lookup(pid: str) -> str:
    return re.sub(r"\s+", "", str(pid or "")).upper()

def _get_suno_title(pid: str) -> Optional[str]:
    return SUNO_TITLES.get(_norm_id_for_lookup(pid))

def translate_site_name(name: str) -> str:
    s = (name or "").strip()
    low = s.lower()
    for key, heb in SITE_NAME_TRANSLATIONS.items():
        if key in low:
            return heb
    return s

def translate_site_name_with_flag(name: str) -> Tuple[str, bool]:
    """
    Return (translated_name, was_translated).
    was_translated=True iff the input matched SITE_NAME_TRANSLATIONS.
    """
    s = (name or "").strip()
    low = s.lower()
    for key, heb in SITE_NAME_TRANSLATIONS.items():
        if key in low:
            return heb, True
    return s, False

# === Header detection & scanning ===

def _is_site_name_header(text: str) -> bool:
    t = text.lower()
    return "site name" in t or ("site" in t and "name" in t)

def _is_coordinates_header(text: str) -> bool:
    t = text.lower()
    return ("coordinates" in t or "coordinate" in t or "coords" in t or "gps" in t)

def _row_has_end_marker(df: pd.DataFrame, row: int) -> bool:
    row_texts = [str(x).strip().lower() for x in df.iloc[row].tolist()]
    return any(any(tok in cell for tok in END_SECTION_TOKENS) for cell in row_texts)

def _find_location_and_coord_headers(df: pd.DataFrame) -> Tuple[Optional[int], Optional[int], Optional[int]]:
    """Find the 'Site name' and 'Coordinates' headers even if they are on different rows."""
    loc_row = loc_col = coord_row = coord_col = None
    for r in range(len(df)):
        for c in range(len(df.columns)):
            txt = str(df.iloc[r, c]).strip()
            if loc_row is None and _is_site_name_header(txt):
                loc_row, loc_col = r, c
            if coord_row is None and _is_coordinates_header(txt):
                coord_row, coord_col = r, c
        if loc_row is not None and coord_row is not None:
            break
    if loc_row is None or coord_row is None:
        return None, None, None
    start_row = max(loc_row, coord_row) + 1
    return loc_col, coord_col, start_row

# === Excel scanning utilities ===

def extract_below_target(df: pd.DataFrame, target: str, max_offset: int = 3) -> str:
    """
    Extract the value below a header.
    - 'Organization' → exact match only
    - 'Date of coordination'/'Date' → prefer any cell containing 'date of coordination';
      else whole-cell 'Date' (avoids matching 'Coordinates')
    - Others → exact-first, then substring
    """
    tgt = target.strip().lower()
    nrows, ncols = len(df), len(df.columns)

    def _scan_from(r: int, c: int) -> str:
        for off in range(1, max_offset + 1):
            nr = r + off
            if nr >= nrows:
                break
            nxt = df.iloc[nr, c]
            if isinstance(nxt, str) and nxt.strip():
                return nxt.strip()
            elif not pd.isna(nxt):
                return str(nxt)
        return "[MISSING DATA BELOW TARGET]"

    if tgt == "organization":
        for r in range(nrows - max_offset):
            for c in range(ncols):
                if str(df.iloc[r, c]).strip().lower() == "organization":
                    return _scan_from(r, c)
        return "[TARGET NOT FOUND]"

    if tgt in ("date", "date of coordination"):
        for r in range(nrows - max_offset):
            for c in range(ncols):
                if "date of coordination" in str(df.iloc[r, c]).strip().lower():
                    return _scan_from(r, c)
        wb = re.compile(r"^\s*date\s*$", re.IGNORECASE)
        for r in range(nrows - max_offset):
            for c in range(ncols):
                if wb.match(str(df.iloc[r, c])):
                    return _scan_from(r, c)
        return "[TARGET NOT FOUND]"

    for r in range(nrows - max_offset):
        for c in range(ncols):
            if str(df.iloc[r, c]).strip().lower() == tgt:
                return _scan_from(r, c)
    for r in range(nrows - max_offset):
        for c in range(ncols):
            if tgt in str(df.iloc[r, c]).strip().lower():
                return _scan_from(r, c)
    return "[TARGET NOT FOUND]"

def extract_hours(df: pd.DataFrame) -> Tuple[str, str]:
    """Return (start_time, end_time) in HH:MM; arrival is shifted by +5h."""
    OFFSET_HOURS = 5

    def _norm_hhmm(s: str) -> str:
        m = re.match(r"^\s*([0-2]?\d):([0-5]\d)\s*$", s)
        if m:
            return f"{int(m.group(1)):02d}:{m.group(2)}"
        return s

    def _coerce_time(value: Any) -> Optional[str]:
        if value is None or (isinstance(value, float) and pd.isna(value)):
            return None
        try:
            if isinstance(value, (int, float)):
                base = datetime(1899, 12, 30) + timedelta(days=float(value))
                return base.strftime("%H:%M")
            s = str(value).strip()
            for fmt in ("%H:%M", "%H:%M:%S"):
                try:
                    return datetime.strptime(s, fmt).strftime("%H:%M")
                except ValueError:
                    pass
            m = re.search(r"([01]?\d|2[0-3]):[0-5]\d", s)
            if m:
                return _norm_hhmm(m.group(0))
            f = float(s)
            base = datetime(1899, 12, 30) + timedelta(days=f)
            return base.strftime("%H:%M")
        except Exception:
            return None

    def _find_header(labels: Sequence[str]) -> Tuple[Optional[int], Optional[int]]:
        labels = [l.lower() for l in labels]
        for r in range(len(df)):
            for c in range(len(df.columns)):
                txt = str(df.iloc[r, c]).strip().lower()
                if any(l in txt for l in labels):
                    return r, c
        return None, None

    def _scan_column_for_times(start_row: Optional[int], col: Optional[int], want: str = "first") -> str:
        if start_row is None or col is None:
            return "[MISSING TIME]"
        times: List[str] = []
        for r in range(start_row + 1, len(df)):
            t = _coerce_time(df.iloc[r, col])
            if t:
                times.append(_norm_hhmm(t))
        if not times:
            return "[MISSING TIME]"
        return times[0] if want == "first" else times[-1]

    dep_row, dep_col = _find_header(["time (hh:mm) of requested departure", "requested departure"])
    arr_row, arr_col = _find_header(["time (hh:mm) of expected arrival", "expected arrival"])
    start_time = _scan_column_for_times(dep_row, dep_col, want="first")
    end_time = _scan_column_for_times(arr_row, arr_col, want="last")
    if end_time != "[MISSING TIME]":
        end_dt = datetime.strptime(end_time, "%H:%M") + timedelta(hours=OFFSET_HOURS)
        end_time = end_dt.strftime("%H:%M")
    return start_time, end_time

def format_and_validate_date(date_value: Any) -> str:
    """Return date as dd.MM.yyyy; accept various common formats or Excel serials."""
    if isinstance(date_value, (int, float)):
        try:
            d = datetime(1899, 12, 30) + timedelta(days=float(date_value))
            return d.strftime("%d.%m.%Y")
        except Exception:
            return "[INVALID DATE]"
    if isinstance(date_value, datetime):
        return date_value.strftime("%d.%m.%Y")
    if isinstance(date_value, str):
        s = _sanitize_text(date_value)
        fmts = [
            "%d.%m.%Y", "%d/%m/%Y",
            "%d.%m.%Y %H:%M", "%d/%m/%Y %H:%M",
            "%Y-%m-%d", "%Y-%m-%d %H:%M:%S",
        ]
        for fmt in fmts:
            try:
                d = datetime.strptime(s, fmt)
                return d.strftime("%d.%m.%Y")
            except ValueError:
                pass
        m = re.search(r"\b(\d{1,2})[./](\d{1,2})[./](\d{4})\b", s)
        if m:
            try:
                d = datetime(int(m.group(3)), int(m.group(2)), int(m.group(1)))
                return d.strftime("%d.%m.%Y")
            except ValueError:
                return "[INVALID DATE]"
    return "[INVALID DATE]"

def latlon_to_utm(lat: float, lon: float) -> str:
    """Convert lat/lon to UTM (zone 36) and shorten northing as in original code."""
    try:
        easting, northing = TRANSFORMER_WGS84_TO_UTM36.transform(lon, lat)
        short_northing = str(int(northing))[1:]  # Remove the leading '3'
        return f"{int(easting)}/{short_northing}"
    except Exception:
        return "[INVALID COORDS]"

# === Coordinates parsing ===

def _coerce_float(val: Any) -> Optional[float]:
    """Loose numeric coercion that never returns nonsense from empty/dirty cells."""
    if val is None or (isinstance(val, float) and pd.isna(val)):
        return None
    s = _sanitize_coord_text(val)
    if s.strip().lower() in _MISSING_TOKENS:
        return None
    if re.match(r"^\s*-?\d+,\d+\s*$", s):
        s = s.replace(",", ".")
    try:
        return float(s)
    except Exception:
        return None

def _inside_region(lat: float, lon: float) -> bool:
    return (LAT_MIN <= lat <= LAT_MAX) and (LON_MIN <= lon <= LON_MAX)

def _parse_lat_lon_from_cell(cell_text: str) -> Optional[Tuple[float, float]]:
    """
    Parse ONLY from a single cell (no whole-row fishing).
    Accepts decimal 'lat, lon' or 'lat lon' or with degrees symbol.
    """
    s = _sanitize_coord_text(cell_text)
    if s.strip().lower() in _MISSING_TOKENS:
        return None

    m = re.search(r"\b([+-]?\d{1,2}(?:\.\d+)?)\s*[,;/\s]\s*([+-]?\d{1,3}(?:\.\d+)?)\b", s)
    if m:
        lat, lon = float(m.group(1)), float(m.group(2))
        if -90 <= lat <= 90 and -180 <= lon <= 180 and _inside_region(lat, lon):
            return lat, lon

    dms_re = re.compile(
        r"([+-]?\d{1,3})\s*[°º]\s*(\d{1,2})\s*['’]?\s*(?:(\d{1,2}(?:\.\d+)?))?\s*(?:[\"”″])?\s*([NSEW])",
        re.IGNORECASE
    )
    parts = list(dms_re.finditer(s))
    if len(parts) >= 2:
        def dms_to_dd(deg_s, min_s, sec_s, hemi):
            deg = float(deg_s); mins = float(min_s); secs = float(sec_s or 0.0)
            dd = abs(deg) + mins/60.0 + secs/3600.0
            if hemi.upper() in ("S", "W") or str(deg_s).startswith("-"):
                dd = -dd
            return dd
        lat_dd = lon_dd = None
        for m in parts[:2]:
            dd = dms_to_dd(*m.groups())
            if m.group(4).upper() in ("N", "S"):
                lat_dd = dd
            else:
                lon_dd = dd
        if lat_dd is not None and lon_dd is not None and -90 <= lat_dd <= 90 and -180 <= lon_dd <= 180 and _inside_region(lat_dd, lon_dd):
            return lat_dd, lon_dd

    return None

def _row_coords_guess(df: pd.DataFrame, r: int, coord_col: Optional[int]) -> Tuple[Optional[Tuple[float, float]], str, str]:
    """
    STRICT: prefer the coordinates cell; then *only* the immediate neighbour pair
    if they both look like lat/lon in range. Never scan whole row (prevents false hits).
    Returns: (latlon|None, reason, raw_cell_used)
    """
    if coord_col is not None:
        raw = str(df.iloc[r, coord_col])
        parsed = _parse_lat_lon_from_cell(raw)
        if parsed:
            return parsed, "parsed_from_coord_col", raw
        if (raw is None) or (str(raw).strip().lower() in _MISSING_TOKENS) or (str(raw).strip() == "") or pd.isna(df.iloc[r, coord_col]):
            return None, "missing_or_empty_coord_cell", str(raw)

    if coord_col is not None and coord_col + 1 < len(df.columns):
        a, b = df.iloc[r, coord_col], df.iloc[r, coord_col + 1]
        la, lb = _coerce_float(a), _coerce_float(b)
        if la is not None and lb is not None:
            candidates = [
                (la, lb, "neighbour_pair_lat_lon"),
                (lb, la, "neighbour_pair_lon_lat")
            ]
            for lat, lon, tag in candidates:
                if -90 <= lat <= 90 and -180 <= lon <= 180 and _inside_region(lat, lon):
                    return (lat, lon), tag, f"{a} | {b}"

    return None, "no_valid_pair_found", ("" if coord_col is None else str(df.iloc[r, coord_col]))

# === Sites & levels ===

def _clean_site_name(name: str) -> str:
    s = (name or "").strip()
    return re.sub(r"^\s*(from|to)\s*:\s*", "", s, flags=re.IGNORECASE)

def extract_sites_for_map(df: pd.DataFrame) -> List[Dict[str, Any]]:
    """Return ordered, de-duplicated sites for mapping; skip rows with missing/invalid coords."""
    location_col = coordinates_col = start_row = None
    for row in range(len(df)):
        for col in range(len(df.columns)):
            cell_value = str(df.iloc[row, col]).strip()
            if _is_site_name_header(cell_value):
                location_col = col; start_row = row + 1
            elif _is_coordinates_header(cell_value):
                coordinates_col = col; start_row = row + 1
    if location_col is None or coordinates_col is None or start_row is None:
        return []

    sites: List[Dict[str, Any]] = []
    last_key: Optional[Tuple[float, float]] = None

    for r in range(start_row, len(df)):
        if _row_has_end_marker(df, r):
            break
        name = df.iloc[r, location_col]
        clean_name = _clean_site_name(name) if isinstance(name, str) else ""
        if not clean_name:
            continue

        parsed, _, _raw = _row_coords_guess(df, r, coordinates_col)
        if not parsed:
            continue  # don't plot missing/invalid ones
        lat, lon = parsed
        key = (round(lat, 6), round(lon, 6))
        if last_key is not None and key == last_key:
            continue
        sites.append({"name": clean_name, "lat": lat, "lon": lon})
        last_key = key

    return sites

def count_coordinates_levels(df: pd.DataFrame) -> List[str]:
    """
    Build the RTL lines for levels using the *same header detection as the map*.
    This prevents choosing the wrong coordinates column and marking everything as missing.
    Also swaps name/coords order when the site name was translated via SITE_NAME_TRANSLATIONS.
    """
    location_col = coordinates_col = start_row = None
    for row in range(len(df)):
        for col in range(len(df.columns)):
            cell_value = str(df.iloc[row, col]).strip()
            if _is_site_name_header(cell_value):
                location_col = col; start_row = row + 1
            elif _is_coordinates_header(cell_value):
                coordinates_col = col; start_row = row + 1
    if location_col is None or coordinates_col is None or start_row is None:
        return ["[MISSING LOCATION/COORDINATES]"]

    entries: List[Dict[str, Any]] = []
    last_key: Optional[Tuple[float, float]] = None

    for r in range(start_row, len(df)):
        if _row_has_end_marker(df, r):
            break
        raw_name = df.iloc[r, location_col]
        clean_name, was_translated = translate_site_name_with_flag(_clean_site_name(raw_name)) if isinstance(raw_name, str) else ("", False)
        if not clean_name:
            continue

        parsed, _, _raw = _row_coords_guess(df, r, coordinates_col)

        if parsed:
            lat, lon = parsed
            key = (round(lat, 6), round(lon, 6))
            if last_key is not None and key == last_key:
                continue
            last_key = key
            utm = latlon_to_utm(lat, lon)
            entries.append({"row": r, "name": clean_name, "parsed": (lat, lon), "utm": utm, "was_translated": was_translated})
        else:
            entries.append({"row": r, "name": clean_name, "parsed": None, "utm": "[MISSING LOCATION/COORDINATES]", "was_translated": was_translated})

    if not entries:
        return ["[MISSING LOCATION/COORDINATES]"]

    valid_idxs = [i for i, e in enumerate(entries) if e["parsed"] is not None]
    label_by_index: Dict[int, str] = {}
    if len(valid_idxs) == 1:
        label_by_index[valid_idxs[0]] = "נקודת התחלה"
    elif len(valid_idxs) >= 2:
        label_by_index[valid_idxs[0]] = "נקודת התחלה"
        label_by_index[valid_idxs[-1]] = "נקודת סיום"
        for i in valid_idxs[1:-1]:
            label_by_index[i] = "נקודת עצירה"

    lines: List[str] = []
    for i, e in enumerate(entries):
        point_type = label_by_index.get(i, "נקודת עצירה")
        loc_span = (f"\u202A{e['name']}\u202C" if not _hebrew(e['name']) else f"\u200F{e['name']}\u200F")
        coords_span = f"\u202A{e['utm']}\u202C"
        label = f"\u200F){point_type}(\u200F"

        if e.get("was_translated", False):
            line = "\u202B" f"{label}\u00A0-\u00A0{loc_span}\u00A0{coords_span}" "\u202C"
        else:
            line = "\u202B" f"{label}\u00A0-\u00A0{coords_span}\u00A0{loc_span}" "\u202C"

        lines.append(line)

    return lines

def _levels_items(df: pd.DataFrame, strip_letter_prefix: bool = False) -> List[str]:
    lines = count_coordinates_levels(df)
    if not strip_letter_prefix:
        return lines
    items: List[str] = []
    for line in lines:
        m = re.match(r"^\s*[\u202A-\u202E]*[^.]+\.\u00A0(.*)$", line)
        items.append(m.group(1) if m else line)
    return items

# === GeoJSON polygons ===
_POLYGONS: List[Any] = []
_GEOJSON_PATH = "south and north polygons.geojson"
if os.path.exists(_GEOJSON_PATH):
    try:
        with open(_GEOJSON_PATH, "r", encoding="utf-8") as f:
            gj = json.load(f)
        feats = gj["features"] if gj.get("type") == "FeatureCollection" else [gj]
        for ft in feats:
            geom = shape(ft["geometry"])
            if geom.geom_type == "Polygon":
                _POLYGONS.append(geom)
            elif geom.geom_type == "MultiPolygon":
                _POLYGONS.extend(list(geom.geoms))
        if len(_POLYGONS) < 2:
            st.warning("GeoJSON loaded but found fewer than 2 polygons.")
    except Exception as e:
        st.warning(f"Failed to read polygons GeoJSON '{_GEOJSON_PATH}': {e}")
else:
    st.info(f"Place your polygons GeoJSON in the app folder as '{_GEOJSON_PATH}'.")

def crossing_cp_from_polys(sites: Sequence[Dict[str, Any]], polys: Sequence[Any]) -> str:
    """Return 'כן' if site points touch 2 or more polygons; otherwise 'לא'."""
    if not sites or not polys:
        return "[MISSING DATA]"
    touched = set()
    for s in sites:
        pt = Point(s["lon"], s["lat"])  # (x=lon, y=lat)
        for i, poly in enumerate(polys):
            try:
                if poly.covers(pt):
                    touched.add(i)
                    if len(touched) >= 2:
                        return "כן"
            except Exception:
                pass
    return "לא"

# === Mapping ===

def generate_sites_map_image(sites: Sequence[Dict[str, Any]], out_path: str = "sites_map.png") -> Optional[str]:
    """Render an OSM basemap with labelled points; robust to single/no points."""
    if not sites:
        return None
    wm = Transformer.from_crs("EPSG:4326", "EPSG:3857", always_xy=True)
    xs: List[float] = []
    ys: List[float] = []
    kept: List[Dict[str, Any]] = []
    for s in sites:
        try:
            x, y = wm.transform(s["lon"], s["lat"])
            if math.isfinite(x) and math.isfinite(y):
                xs.append(x); ys.append(y); kept.append(s)
        except Exception:
            continue
    if not xs:
        return None
    fig, ax = plt.subplots(figsize=(6.5, 5), dpi=200)
    ax.scatter(xs, ys, s=30, zorder=3)
    for (x, y), s in zip(zip(xs, ys), kept):
        ax.annotate(
            f"{s['name']}\n{s['lat']:.6f}, {s['lon']:.6f}",
            (x, y), xytext=(0, 10), textcoords="offset points",
            ha="center", va="bottom", fontsize=9, zorder=4,
            bbox=dict(facecolor="white", alpha=0.7, edgecolor="none", pad=1.5),
        )
    if len(xs) == 1:
        pad_x = pad_y = 500
        ax.set_xlim(xs[0] - pad_x, xs[0] + pad_x)
        ax.set_ylim(ys[0] - pad_y, ys[0] + pad_y)
    else:
        pad_x = max(200, (max(xs) - min(xs)) * 0.15)
        pad_y = max(200, (max(ys) - min(ys)) * 0.15)
        ax.set_xlim(min(xs) - pad_x, max(xs) + pad_x)
        ax.set_ylim(min(ys) - pad_y, max(ys) + pad_y)
    try:
        cx.add_basemap(ax, source=cx.providers.OpenStreetMap.Mapnik, crs="EPSG:3857", attribution=False)
        ax.set_axis_off()
    except Exception:
        ax.grid(True, linestyle=":", linewidth=0.5)
    fig.tight_layout()
    fig.savefig(out_path, dpi=200, bbox_inches="tight", pad_inches=0.05)
    plt.close(fig)
    return out_path

# === Word helpers ===

def set_font(cell, font_name: str = "David", font_size: int = 12, align_center: bool = True, vcenter: bool = True) -> None:
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        if align_center:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

def _find_paragraph_with_text(doc: Document, token: str) -> Optional[Paragraph]:
    for para in doc.paragraphs:
        if token in para.text:
            return para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if token in para.text:
                        return para
    return None

def _set_para_text_preserve_format(para: Paragraph, text: str) -> None:
    runs = list(para.runs)
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r.text = ""
    else:
        para.add_run(text)

def _replace_token_with_numbered_list(doc_path: str, token: str, items: Sequence[str]) -> str:
    doc = Document(doc_path)
    model_para = _find_paragraph_with_text(doc, token)
    if not model_para:
        doc.save(doc_path)
        return doc_path
    parent = model_para._p.getparent()
    insert_at = parent.index(model_para._p)
    if items:
        _set_para_text_preserve_format(model_para, items[0])
    for text in items[1:]:
        new_p = deepcopy(model_para._p)
        insert_at += 1
        parent.insert(insert_at, new_p)
        new_para = Paragraph(new_p, model_para._parent)
        _set_para_text_preserve_format(new_para, text)
    doc.save(doc_path)
    return doc_path

def _clear_data_rows_keep_model(table, model_row_idx: int = 1) -> None:
    while len(table.rows) > model_row_idx + 1:
        table._element.remove(table.rows[model_row_idx + 1]._element)

def _clone_row_with_numbering(table, model_row_idx: int = 1):
    tbl = table._tbl
    tr = tbl.tr_lst[model_row_idx]
    new_tr = deepcopy(tr)
    tbl.append(new_tr)
    return table.rows[-1]

# === Personnel & Vehicles extraction ===

def check_white_wheels(ids: Iterable[str]) -> str:
    for id_value in ids:
        if isinstance(id_value, str) and id_value.strip():
            first_digit = id_value.strip()[0]
            if first_digit not in {"3", "4", "8", "9"}:
                return "גלגלים לבנים"
    return ""

def translate_personnel_role(text: Any) -> str:
    if not isinstance(text, str) or not text.strip():
        return "[MISSING DATA]"
    original_text = text.strip()
    lowered = original_text.lower()
    main_role = ""
    notes: List[str] = []
    for eng_term, heb_term in PERSONNEL_TRANSLATIONS.items():
        if eng_term in lowered:
            if "(" in heb_term:
                notes.append(heb_term)
            else:
                main_role = heb_term
    if main_role:
        return f"\u202B{main_role} {' '.join(notes)}\u202C" if notes else f"\u202B{main_role}\u202C"
    return f"\u202B{' '.join(notes)}\u202C" if notes else original_text

def _combine_full_name(first_name: str, last_name: str) -> str:
    fn = (first_name or "").strip()
    ln = (last_name or "").strip()
    if not fn and not ln:
        return "[MISSING DATA]"
    if fn and ln and fn.lower() == ln.lower():
        return fn
    if fn and ln:
        if fn.lower() in ln.lower() and len(ln) >= len(fn):
            return ln
        if ln.lower() in fn.lower() and len(fn) >= len(ln):
            return fn
        parts = (fn + " " + ln).split()
        seen: set[str] = set()
        deduped: List[str] = []
        for p in parts:
            key = p.lower()
            if key not in seen:
                seen.add(key)
                deduped.append(p)
        return " ".join(deduped)
    return fn or ln

def _clean_id_text(s: str) -> str:
    if not isinstance(s, str):
        s = "" if s is None else str(s)
    s = s.strip()
    m = re.search(r"\bID\b\s*[:\-]?\s*([A-Z0-9]+)\b", s, flags=re.IGNORECASE)
    return m.group(1) if m else s

def extract_personnel_table_data(df: pd.DataFrame) -> Tuple[List[Dict[str, str]], str]:
    required_headers: Dict[str, Optional[int]] = {
        "first name": None,
        "last name": None,
        "id": None,
        "mission of the specific personnel": None,
    }
    personnel_start_row: Optional[int] = None
    for row in range(len(df)):
        row_values = [str(cell).strip().lower() for cell in df.iloc[row]]
        if all(any(key in cell for cell in row_values) for key in required_headers.keys()):
            personnel_start_row = row
            for col in range(len(df.columns)):
                cell_value = str(df.iloc[row, col]).strip().lower()
                for key in list(required_headers.keys()):
                    if key in cell_value and required_headers[key] is None:
                        required_headers[key] = col
            break
    if personnel_start_row is None or any(v is None for v in required_headers.values()):
        return [], "[MISSING PERSONNEL DATA]"

    extracted_data: List[Dict[str, str]] = []
    id_list: List[str] = []
    empty_row_count = 0
    max_empty_rows_allowed = 1

    def get_cell(row: int, col_idx: Optional[int]) -> str:
        if col_idx is None:
            return ""
        val = df.iloc[row, col_idx]
        if pd.isna(val) or str(val).strip() == "":
            return ""
        if isinstance(val, float):
            return str(int(val))
        return str(val).strip()

    for row in range(personnel_start_row + 1, len(df)):
        first_name = get_cell(row, required_headers["first name"])  # type: ignore[arg-type]
        last_name = get_cell(row, required_headers["last name"])   # type: ignore[arg-type]
        pid = _clean_id_text(get_cell(row, required_headers["id"])) # type: ignore[arg-type]
        mission_txt = get_cell(row, required_headers["mission of the specific personnel"])  # type: ignore[arg-type]

        is_empty_row = not any([first_name, last_name, pid, mission_txt])
        if is_empty_row:
            empty_row_count += 1
            if empty_row_count > max_empty_rows_allowed:
                break
            continue
        empty_row_count = 0

        full_name = _combine_full_name(first_name, last_name)
        row_data = {
            "name": full_name,
            "id": pid if pid else "[MISSING DATA]",
            "notes": mission_txt if mission_txt else "[MISSING DATA]",
        }
        if pid:
            id_list.append(pid)
        extracted_data.append(row_data)

    white_wheels_text = check_white_wheels(id_list)
    return extracted_data, white_wheels_text

def translate_vehicle_type(vehicle_type: Any) -> str:
    if not isinstance(vehicle_type, str) or not vehicle_type.strip():
        return "[MISSING DATA]"
    normalized_type = vehicle_type.lower().strip()
    for eng_type, hebrew_translation in VEHICLE_TRANSLATIONS.items():
        if eng_type in normalized_type:
            brand_name = normalized_type.replace(eng_type, "").strip().title()
            return f"{brand_name}\u00A0{hebrew_translation}" if brand_name else hebrew_translation
    return vehicle_type

def extract_vehicle_table_data(df: pd.DataFrame) -> Tuple[List[Dict[str, str]], str]:
    target_headers: Dict[str, Optional[int]] = {"vehicle type": None, "plate number": None}
    vehicle_start_row: Optional[int] = None
    empty_row_count = 0
    max_empty_rows_allowed = 1

    for row in range(len(df)):
        for col in range(len(df.columns)):
            cell_value = str(df.iloc[row, col]).strip().lower()
            for key in list(target_headers.keys()):
                if key in cell_value and target_headers[key] is None:
                    target_headers[key] = col
                    if vehicle_start_row is None:
                        vehicle_start_row = row
        if all(v is not None for v in target_headers.values()):
            break
    if any(v is None for v in target_headers.values()) or vehicle_start_row is None:
        return [], "[MISSING VEHICLE DATA]"

    extracted_vehicles: List[Dict[str, str]] = []
    for row in range(vehicle_start_row + 1, len(df)):
        is_empty_row = True
        row_data: Dict[str, str] = {}
        for key, col_index in target_headers.items():
            val = df.iloc[row, col_index]  # type: ignore[index]
            if pd.isna(val) or str(val).strip() == "":
                cell_value = "[MISSING DATA]"
            elif isinstance(val, float):
                cell_value = str(int(val))
            else:
                cell_value = str(val).strip()
            if cell_value != "[MISSING DATA]":
                is_empty_row = False
            row_data[key] = cell_value
        if is_empty_row:
            empty_row_count += 1
            if empty_row_count > max_empty_rows_allowed:
                break
        else:
            empty_row_count = 0
            extracted_vehicles.append(row_data)
    return extracted_vehicles, ""

# === Word table updaters ===

def _fill_numbered_table(table, rows: Sequence[Sequence[str]]) -> None:
    """Assumes table row[0] is header and row[1] is a model row with numbering in the first cell."""
    _clear_data_rows_keep_model(table, model_row_idx=1)
    if not rows:
        return
    row0_cells = table.rows[1].cells
    for i, text in enumerate(rows[0], start=1):  # skip number col at index 0
        row0_cells[i].text = text
    for cell in row0_cells:
        set_font(cell)
    for data in rows[1:]:
        new_row = _clone_row_with_numbering(table, model_row_idx=1).cells
        for i, text in enumerate(data, start=1):
            new_row[i].text = text
        for cell in new_row:
            set_font(cell)

def update_personnel_table(doc_path: str, extracted_data: List[Dict[str, str]], white_wheels_text: str) -> str:
    doc = Document(doc_path)
    if not doc.tables:
        return "[NO TABLE FOUND IN DOCUMENT]"
    table = doc.tables[0]

    rows: List[Tuple[str, str, str]] = []
    bold_indices: List[int] = []
    for idx, row in enumerate(extracted_data):
        role_txt = translate_personnel_role(row.get("notes", ""))
        name_txt = row.get("name", "[MISSING DATA]")
        id_txt = row.get("id", "[MISSING DATA]")
        suno_title = _get_suno_title(id_txt)
        if suno_title:
            name_txt = f"{name_txt} - {suno_title}"
            bold_indices.append(idx)
        rows.append((role_txt, name_txt, id_txt))

    _fill_numbered_table(table, rows)

    for idx in bold_indices:
        row = table.rows[idx + 1]  # +1 for header
        for ci in (2, 3):
            for p in row.cells[ci].paragraphs:
                for r in p.runs:
                    r.font.bold = True

    if white_wheels_text.strip():
        for para in doc.paragraphs:
            if "{ white_wheels }" in para.text:
                para.text = para.text.replace("{ white_wheels }", white_wheels_text)
    else:
        for shape_el in doc.element.xpath("//w:drawing"):
            parent = shape_el.getparent()
            if parent is not None:
                parent.remove(shape_el)
                break

    doc.save("updated_document.docx")
    return "updated_document.docx"

def update_vehicle_table(doc_path: str, extracted_vehicles: List[Dict[str, str]]) -> str:
    doc = Document(doc_path)
    if len(doc.tables) < 2:
        return "[SECOND TABLE NOT FOUND]"
    table = doc.tables[1]
    rows: List[Tuple[str, str]] = []
    for r in extracted_vehicles:
        rows.append((translate_vehicle_type(r.get("vehicle type", "[MISSING DATA]")), r.get("plate number", "[MISSING DATA]")))
    _fill_numbered_table(table, rows)
    doc.save("updated_document.docx")
    return "updated_document.docx"

# === Highlighting ===

def update_highlighting(doc_path: str, start_date: str, equipment_list: str) -> str:
    from docx.enum.text import WD_COLOR_INDEX
    doc = Document(doc_path)

    def _clone_rpr(dst_run, src_run):
        rPr = src_run._element.rPr
        if rPr is not None:
            dst_run._element.insert(0, deepcopy(rPr))
        else:
            dst_run.font.name = src_run.font.name
            if src_run.font.size:
                dst_run.font.size = src_run.font.size
            dst_run.font.bold = src_run.font.bold
            dst_run.font.italic = src_run.font.italic
            dst_run.font.underline = src_run.font.underline
            try:
                if src_run.font.color and src_run.font.color.rgb:
                    dst_run.font.color.rgb = src_run.font.color.rgb
            except Exception:
                pass

    def _highlight_exact(paragraph, token: str, color):
        if not token or token not in paragraph.text:
            return
        i = 0
        while i < len(paragraph.runs):
            run = paragraph.runs[i]
            txt = run.text or ""
            pos = txt.find(token)
            if pos == -1:
                i += 1
                continue
            before, after = txt[:pos], txt[pos + len(token):]
            run.text = before
            tok_run = paragraph.add_run(token)
            _clone_rpr(tok_run, run)
            tok_run.font.highlight_color = color
            run._element.addnext(tok_run._element)
            if after:
                after_run = paragraph.add_run(after)
                _clone_rpr(after_run, run)
                tok_run._element.addnext(after_run._element)
                i += 2
            else:
                i += 2

    tomorrow = (datetime.now() + timedelta(days=1)).strftime("%d.%m.%Y")
    if isinstance(start_date, str) and start_date and start_date != tomorrow:
        for para in list(doc.paragraphs):
            if start_date in para.text:
                _highlight_exact(para, start_date, WD_COLOR_INDEX.RED)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if start_date in para.text:
                            _highlight_exact(para, start_date, WD_COLOR_INDEX.RED)

    if isinstance(equipment_list, str) and equipment_list.strip().startswith("יש"):
        token = "יש"
        for para in list(doc.paragraphs):
            if token in para.text:
                _highlight_exact(para, token, WD_COLOR_INDEX.YELLOW)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if token in para.text:
                            _highlight_exact(para, token, WD_COLOR_INDEX.YELLOW)

    missing_data_errors = [
        "[MISSING DATA]", "[MISSING TIME]", "[MISSING LOCATION/COORDINATES]",
        "[MISSING VEHICLE DATA]", "[TARGET NOT FOUND]", "[INVALID DATE]", "[INVALID COORDS]",
    ]
    for para in list(doc.paragraphs):
        txt = para.text
        if not txt:
            continue
        for err in missing_data_errors:
            if err in txt:
                _highlight_exact(para, err, WD_COLOR_INDEX.RED)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    txt = para.text
                    if not txt:
                        continue
                    for err in missing_data_errors:
                        if err in txt:
                            _highlight_exact(para, err, WD_COLOR_INDEX.RED)

    doc.save("highlighted_document.docx")
    return "highlighted_document.docx"

# === Streamlit App ===

def run_app() -> None:
    st.title("Coordinations App")
    uploaded_excel = st.file_uploader("Upload your Excel file (.xlsx, .xlsm)", type=["xlsx", "xlsm"])
    user_name = st.text_input("Enter your name:")

    if not (uploaded_excel and user_name):
        return

    # Read Excel
    try:
        df = pd.read_excel(uploaded_excel, engine="openpyxl", header=None)
        df = df.applymap(lambda x: x.strip() if isinstance(x, str) else x)
    except Exception as e:
        st.error(f"Error reading Excel file: {e}")
        st.stop()
    st.dataframe(df)

    # Load template
    try:
        doc_tpl = DocxTemplate(TEMPLATE_FILE)
    except Exception as e:
        st.error(f"Error loading Word template: {e}")
        st.stop()

    # Extract fields
    organization_name = extract_below_target(df, "ORGANIZATION")
    raw_start_date = extract_below_target(df, "Date of coordination")
    coordination_header = extract_below_target(df, "Objective of the mission")
    start_time, end_time = extract_hours(df)
    start_date = format_and_validate_date(raw_start_date)
    level_items = _levels_items(df, strip_letter_prefix=True)

    # Sites & map
    sites = extract_sites_for_map(df)
    crossing_cp = crossing_cp_from_polys(sites, _POLYGONS)
    sites_map_img = ""
    if sites:
        img_path = generate_sites_map_image(sites, out_path="sites_map.png")
        if img_path:
            sites_map_img = InlineImage(doc_tpl, img_path, width=Cm(12))

    mission = extract_below_target(df, "Detailed execution")
    equipment_list = extract_equipment_status(df)

    table_data, white_wheels_text = extract_personnel_table_data(df)
    vehicle_table_data, _veh_status = extract_vehicle_table_data(df)

    data_dict = {
        "organization_name": organization_name,
        "start_date": start_date,
        "end_date": start_date,
        "start_time": start_time,
        "end_time": end_time,
        "levels": "[[LEVELS_LIST]]",  # placeholder; replaced with numbered list
        "white_wheels": white_wheels_text,
        "soldier_name": user_name,
        "coordination_header": coordination_header,
        "mission": mission,
        "equipment_list": equipment_list,
        "sites_map": sites_map_img,
        "crossing_cp": crossing_cp,
    }

    # Render, save, replace levels, fill tables, highlight, and expose download
    doc_tpl.render(data_dict)
    doc_tpl.save("processed_document.docx")

    _replace_token_with_numbered_list("processed_document.docx", "[[LEVELS_LIST]]", level_items)

    final_doc = update_personnel_table("processed_document.docx", table_data, white_wheels_text)
    final_doc = update_vehicle_table(final_doc, vehicle_table_data)
    final_doc = update_highlighting(final_doc, start_date, equipment_list)

    with open(final_doc, "rb") as f:
        output_stream = BytesIO(f.read())

    st.download_button(
        label="Download Processed Word Document",
        data=output_stream,
        file_name="processed_document.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# === Equipment scan (kept identical in behaviour) ===

def extract_equipment_status(df: pd.DataFrame) -> str:
    """
    Scan rows between 'Equipment' and two rows above 'NOTES'/mission text to see if any cell is 'yes'.
    Returns 'יש ציוד' if found, else 'אין ציוד'.
    """
    equipment_row = equipment_col = None
    notes_row: Optional[int] = None

    # Find 'Equipment'
    for row in range(len(df)):
        for col in range(len(df.columns)):
            cell_value = str(df.iloc[row, col]).strip().lower()
            if "equipment" in cell_value:
                equipment_row, equipment_col = row, col
                break
        if equipment_row is not None:
            break

    # Find 'mission of the specific personnel' (serves as NOTES delimiter)
    for row in range(len(df)):
        for col in range(len(df.columns)):
            cell_value = str(df.iloc[row, col]).strip().lower()
            if "mission of the specific personnel" in cell_value:
                notes_row = row
                break
        if notes_row is not None:
            break

    if equipment_row is None or notes_row is None:
        return "[MISSING EQUIPMENT OR NOTES]"

    for row in range(equipment_row + 1, notes_row - 1):
        for col in range(len(df.columns)):
            if str(df.iloc[row, col]).strip().lower() == "yes":
                return "יש ציוד"
    return "אין ציוד"

# === Entrypoint for Streamlit ===
if __name__ == "__main__":
    run_app()
